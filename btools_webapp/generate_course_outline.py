import sys
import os
import re
import json
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

STRICT_FONT_NAME = "Sarabun"

def clean_bullet_text(text):
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    # ไม่ลบสัญลักษณ์ใดๆ ทั้งสิ้น เพื่อคงรูปแบบดั้งเดิม (ตัวเลข, ขีด, ลูกศร) ตามที่ AI ส่งมา
    return text.strip()

def add_smart_bullet(p, text, font_size=10, bold=False):
    if text is None:
        return None
    if not isinstance(text, str):
        text = str(text)
    
    raw_text = text.strip()
    # 1. เช็คว่าขึ้นต้นด้วยตัวเลข (เช่น 1., a.) หรือลูกศร (->, ➢) หรือไม่
    is_numbered_or_arrow = bool(re.match(r'^(\d+[\.\)]|[a-zA-Z][\.\)]|->|=>|➢|➔|▪|►)\s', raw_text))
    
    # 2. ลบสัญลักษณ์ Bullet พื้นฐาน (จุด, ขีด) ทิ้งเพื่อไม่ให้ซ้อนกัน
    clean_text = re.sub(r'^[\s•\-\*✓👤]+', '', raw_text).strip()
    
    # 3. ถ้าไม่ใช่ตัวเลข/ลูกศร ให้ใส่ Native Word Bullet (w:numPr)
    if not is_numbered_or_arrow and clean_text:
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), '1')
        numPr.append(ilvl)
        numPr.append(numId)
        pPr.append(numPr)
        
    final_text = clean_text if not is_numbered_or_arrow else raw_text
    
    # รองรับการทำตัวหนาด้วย Markdown (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', final_text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = p.add_run(part[2:-2])
            run.font.name = STRICT_FONT_NAME
            run.font.size = Pt(font_size)
            run.bold = True
        else:
            run = p.add_run(part)
            run.font.name = STRICT_FONT_NAME
            run.font.size = Pt(font_size)
            run.bold = bold
    return p

DEFAULT_TEMPLATE = os.path.join(os.path.dirname(__file__), "templates", "template.docx")

def sanitize_pgmar(doc):
    for sec in doc.sections:
        pgMar = sec._sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            for k in list(pgMar.attrib.keys()):
                val = pgMar.attrib[k]
                try:
                    pgMar.set(k, str(int(round(float(val)))))
                except:
                    pass

def apply_table_borders(table):
    """กำหนดเส้นขอบตารางแบบ 1pt สีดำ ชัดเจนทุกช่อง"""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for b_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{b_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '8')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b_name in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{b_name}')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '8')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), '000000')
                tcBorders.append(b)
            tcPr.append(tcBorders)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_col_widths(table, col_widths_in_inches):
    col_dxas = [int(w * 1440) for w in col_widths_in_inches]
    total_dxa = sum(col_dxas)
    tblPr = table._tbl.tblPr
    
    # 1. Set Table Width (tblW)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total_dxa))
    tblW.set(qn('w:type'), 'dxa')

    # 2. Set Table Left Indent (tblInd) to 0 dxa so left border is perfectly flush with left margin
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')

    # 3. Set Table Grid (tblGrid)
    tblGrid = table._tbl.tblGrid
    if tblGrid is not None:
        tblGrid.getparent().remove(tblGrid)
    
    new_tblGrid = OxmlElement('w:tblGrid')
    for dxa in col_dxas:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(dxa))
        new_tblGrid.append(gc)
    
    table._tbl.insert(table._tbl.index(tblPr) + 1, new_tblGrid)

    # 4. Set Column and Cell widths
    for i, col in enumerate(table.columns):
        col.width = Inches(col_widths_in_inches[i])

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(col_widths_in_inches[i])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(col_dxas[i]))
            tcW.set(qn('w:type'), 'dxa')

def add_blank_line(doc, font_size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run()
    run.font.name = STRICT_FONT_NAME
    run.font.size = Pt(font_size)
    return p

def add_heading(doc, text, font_size=12, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = STRICT_FONT_NAME
    run.font.size = Pt(font_size)
    run.bold = bold
    return p

def add_rationale_p(doc, text, font_size=10):
    """ย่อหน้าเนื้อหาทั่วไป (ไม่มี bullet, margin ปกติ) และรองรับ **ตัวหนา**"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    
    parts = re.split(r'(\*\*.*?\*\*)', str(text))
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            run = p.add_run(part[2:-2])
            run.font.name = STRICT_FONT_NAME
            run.font.size = Pt(font_size)
            run.bold = True
        else:
            run = p.add_run(part)
            run.font.name = STRICT_FONT_NAME
            run.font.size = Pt(font_size)
    return p

def add_bullet_p(doc, text, font_size=10, bold=False, left_indent_in=0.5, hanging_in=0.25):
    """สร้างย่อหน้าแบบ Native Word Bullet (<w:numPr>) พร้อมการจัดระยะเยื้องแบบกำหนดได้"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.left_indent = Inches(left_indent_in)
    p.paragraph_format.first_line_indent = Inches(-hanging_in)

    run = add_smart_bullet(p, text, font_size=font_size, bold=bold)
    return p

def is_intro_sentence(text):
    if not text or not isinstance(text, str):
        return False
    t = text.strip()
    intro_keywords = ["เพื่อให้", "เมื่อจบ", "เมื่อสิ้นสุด", "วัตถุประสงค์", "ผู้เข้าอบรมจะสามารถ", "ผู้เข้าอบรมสามารถ"]
    return any(kw in t for kw in intro_keywords) or t.endswith(":") or t.endswith("ดังนี้") or t.endswith("สามารถ")

def add_subhead_p(doc, text, font_size=10, left_indent_in=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.left_indent = Inches(left_indent_in)
    run = p.add_run(text.strip())
    run.font.name = STRICT_FONT_NAME
    run.font.size = Pt(font_size)
    run.bold = False
    return p

def fix_footer_page_number(doc):
    """จัดวางเลขหน้าใน Footer ให้อยู่ขวาล่างชิดขอบขวา อย่างสวยงาม ด้วย Right Tab Stop 6.92 in"""
    for sec in doc.sections:
        footer = sec.footer
        for p in footer.paragraphs:
            if 'www.btoolstraining.com' in p.text or any('PAGE' in (n.text or n.attrib.get(qn('w:instr'), '')) for n in p._p.iter()):
                pPr = p._p.get_or_add_pPr()
                tabs = pPr.find(qn('w:tabs'))
                if tabs is not None:
                    pPr.remove(tabs)
                tabs = OxmlElement('w:tabs')
                tab = OxmlElement('w:tab')
                tab.set(qn('w:val'), 'right')
                tab.set(qn('w:pos'), '9965') # 6.92 inches * 1440 = 9965 dxa
                tabs.append(tab)
                pPr.append(tabs)

                drawings = p._p.xpath('.//w:drawing')

                for child in list(p._p):
                    if child.tag != qn('w:pPr'):
                        p._p.remove(child)

                # Left run: www.btoolstraining.com (12pt Bold)
                r1 = OxmlElement('w:r')
                rPr1 = OxmlElement('w:rPr')
                rFont1 = OxmlElement('w:rFonts')
                rFont1.set(qn('w:ascii'), STRICT_FONT_NAME)
                rFont1.set(qn('w:hAnsi'), STRICT_FONT_NAME)
                rPr1.append(rFont1)
                sz1 = OxmlElement('w:sz')
                sz1.set(qn('w:val'), '24') # 12pt
                rPr1.append(sz1)
                b1 = OxmlElement('w:b') # Bold
                rPr1.append(b1)
                color1 = OxmlElement('w:color')
                color1.set(qn('w:val'), '434343')
                rPr1.append(color1)
                r1.append(rPr1)
                t1 = OxmlElement('w:t')
                t1.text = 'www.btoolstraining.com'
                r1.append(t1)
                p._p.append(r1)

                # Single tab to right margin
                r_tab = OxmlElement('w:r')
                r_tab.append(OxmlElement('w:tab'))
                p._p.append(r_tab)

                # Dynamic PAGE Field
                fldSimple = OxmlElement('w:fldSimple')
                fldSimple.set(qn('w:instr'), 'PAGE')
                r2 = OxmlElement('w:r')
                rPr2 = OxmlElement('w:rPr')
                rFont2 = OxmlElement('w:rFonts')
                rFont2.set(qn('w:ascii'), STRICT_FONT_NAME)
                rFont2.set(qn('w:hAnsi'), STRICT_FONT_NAME)
                rPr2.append(rFont2)
                sz2 = OxmlElement('w:sz')
                sz2.set(qn('w:val'), '20') # 10pt
                rPr2.append(sz2)
                color2 = OxmlElement('w:color')
                color2.set(qn('w:val'), '434343')
                rPr2.append(color2)
                r2.append(rPr2)
                fldSimple.append(r2)
                p._p.append(fldSimple)

                for dwg in drawings:
                    r_dwg = OxmlElement('w:r')
                    r_dwg.append(dwg)
                    p._p.append(r_dwg)

def render_workshop_cell(cell, w_data):
    if not w_data:
        return
    p_w = cell.add_paragraph()
    p_w.paragraph_format.space_before = Pt(3)
    p_w.paragraph_format.space_after = Pt(2)
    p_w.paragraph_format.line_spacing = 1.35
    
    if isinstance(w_data, dict):
        title = w_data.get("title") or "Workshop"
        bullets = w_data.get("bullets") or w_data.get("topics") or []
        run_b = p_w.add_run(f"Workshop ({title}): " if not title.startswith("Workshop") else f"{title}: ")
        run_b.font.name = STRICT_FONT_NAME
        run_b.font.size = Pt(10)
        run_b.bold = True
        
        for b_item in bullets:
            p_b = cell.add_paragraph()
            p_b.paragraph_format.space_before = Pt(0)
            p_b.paragraph_format.space_after = Pt(2)
            p_b.paragraph_format.line_spacing = 1.35
            p_b.paragraph_format.left_indent = Inches(0.28)
            p_b.paragraph_format.first_line_indent = Inches(-0.12)
            
            pPr = p_b._p.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')
            numPr.append(ilvl)
            numPr.append(numId)
            pPr.append(numPr)
            
            run_n = p_b.add_run(clean_bullet_text(b_item))
            run_n.font.name = STRICT_FONT_NAME
            run_n.font.size = Pt(10)
    else:
        w_text = str(w_data).strip()
        lines = [line.strip() for line in w_text.split('\n') if line.strip()]
        header_line = lines[0] if lines else w_text
        bullet_lines = lines[1:] if len(lines) > 1 else []
        
        if header_line.startswith("Workshop") or header_line.startswith("กิจกรรม"):
            if ":" in header_line:
                parts = header_line.split(":", 1)
                bold_prefix = f"{parts[0].strip()}: "
                desc_text = parts[1].strip()
            else:
                bold_prefix = f"{header_line}: "
                desc_text = ""
        else:
            if ":" in header_line:
                parts = header_line.split(":", 1)
                bold_prefix = f"Workshop ({parts[0].strip()}): "
                desc_text = parts[1].strip()
            else:
                bold_prefix = "Workshop : "
                desc_text = header_line
                
        run_b = p_w.add_run(bold_prefix)
        run_b.font.name = STRICT_FONT_NAME
        run_b.font.size = Pt(10)
        run_b.bold = True
        
        if desc_text:
            run_n = p_w.add_run(desc_text)
            run_n.font.name = STRICT_FONT_NAME
            run_n.font.size = Pt(10)
            
        for b_item in bullet_lines:
            p_b = cell.add_paragraph()
            p_b.paragraph_format.space_before = Pt(0)
            p_b.paragraph_format.space_after = Pt(2)
            p_b.paragraph_format.line_spacing = 1.35
            p_b.paragraph_format.left_indent = Inches(0.28)
            p_b.paragraph_format.first_line_indent = Inches(-0.12)
            
            pPr = p_b._p.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')
            numPr.append(ilvl)
            numPr.append(numId)
            pPr.append(numPr)
            
            run_n = p_b.add_run(clean_bullet_text(b_item))
            run_n.font.name = STRICT_FONT_NAME
            run_n.font.size = Pt(10)

def generate_doc(data, output_path, template_path=None):
    global _CURRENT_AI_MODEL
    _CURRENT_AI_MODEL = data.get("_ai_model_used", "")
    
    if not template_path or not os.path.exists(template_path):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        rel_template = os.path.abspath(os.path.join(script_dir, "..", "templates", "template.docx"))
        if os.path.exists(rel_template):
            template_path = rel_template
        elif os.path.exists(DEFAULT_TEMPLATE):
            template_path = DEFAULT_TEMPLATE
        elif os.path.exists(r"D:\Course outline\00_Reference_Template\B Tools_หลักสูตร Data Analysis for Better Results.docx"):
            template_path = r"D:\Course outline\00_Reference_Template\B Tools_หลักสูตร Data Analysis for Better Results.docx"
            
    if template_path and os.path.exists(template_path):
        doc = docx.Document(template_path)
        sanitize_pgmar(doc)
        for p in list(doc.paragraphs):
            p._p.getparent().remove(p._p)
        for t in list(doc.tables):
            t._element.getparent().remove(t._element)
    else:
        doc = docx.Document()
        for section in doc.sections:
            section.top_margin = Inches(0.787)
            section.bottom_margin = Inches(0.59)
            section.left_margin = Inches(0.787)
            section.right_margin = Inches(0.787)

    fix_footer_page_number(doc)
    sanitize_pgmar(doc)

    # Title TH (H1: 13pt Bold)
    if data.get("course_title_th"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"หลักสูตร: {data['course_title_th']}")
        run.font.name = STRICT_FONT_NAME
        run.font.size = Pt(13)
        run.bold = True

    # Title EN
    if data.get("course_title_en"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(data['course_title_en'])
        run.font.name = STRICT_FONT_NAME
        run.font.size = Pt(13)
        run.bold = True

    add_blank_line(doc)

    # Dynamic Section Ordering (เรียงลำดับ H2 ตามต้นฉบับ 100%)
    sections_order = data.get("sections_order")
    if not sections_order:
        sections_order = [
            "rationale", "objectives", "topics", "agenda",
            "learning_methods", "target_audience", "duration",
            "frameworks", "workshop_activities", "teaching_style",
            "expected_outcomes", "remarks"
        ]

    for sec_key in sections_order:
        if sec_key == "rationale" and data.get("rationale"):
            add_heading(doc, data.get("rationale_title", "หลักการและเหตุผล"), font_size=12)
            for item in data["rationale"]:
                if isinstance(item, dict):
                    if item.get("text"):
                        add_rationale_p(doc, item["text"])
                    if item.get("bullets"):
                        for b in item["bullets"]:
                            add_bullet_p(doc, b)
                        add_blank_line(doc)
                    elif item.get("text"):
                        add_blank_line(doc)
                elif isinstance(item, str):
                    add_rationale_p(doc, item)
                    add_blank_line(doc)

        elif sec_key == "objectives" and data.get("objectives"):
            add_heading(doc, data.get("objectives_title", "วัตถุประสงค์ของหลักสูตร"), font_size=12)
            if data.get("objectives_subhead"):
                add_subhead_p(doc, data["objectives_subhead"])
            for item in data["objectives"]:
                if isinstance(item, dict):
                    t_title = item.get("title") or item.get("text") or ""
                    if t_title:
                        add_subhead_p(doc, t_title)
                    sub_list = item.get("sub_bullets") or item.get("bullets") or []
                    for sub in sub_list:
                        add_bullet_p(doc, sub, left_indent_in=0.5, hanging_in=0.25)
                elif isinstance(item, str):
                    if is_intro_sentence(item):
                        add_subhead_p(doc, item)
                    else:
                        add_bullet_p(doc, item, left_indent_in=0.5, hanging_in=0.25)
            add_blank_line(doc)

        elif sec_key == "topics" and data.get("topics"):
            add_heading(doc, data.get("topics_title", data.get("agenda_title", "หัวข้อการเรียนรู้")), font_size=12)
            for item in data["topics"]:
                if isinstance(item, dict):
                    if item.get("title"):
                        add_bullet_p(doc, item["title"], left_indent_in=0.5, hanging_in=0.25)
                    if item.get("sub_bullets"):
                        for sub in item["sub_bullets"]:
                            add_bullet_p(doc, sub, left_indent_in=0.75, hanging_in=0.25)
                elif isinstance(item, str):
                    add_bullet_p(doc, item, left_indent_in=0.5, hanging_in=0.25)
            add_blank_line(doc)

        elif sec_key == "agenda" and data.get("agenda"):
            add_heading(doc, data.get("agenda_title", "Agenda"), font_size=12)
            has_time = any((item.get("time") or "").strip() for item in data["agenda"] if isinstance(item, dict))
            if has_time:
                table = doc.add_table(rows=1, cols=2)
                try:
                    table.style = 'Table Grid'
                except Exception:
                    pass
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False

                hdr_cells = table.rows[0].cells
                hdr_col1_name = data.get("agenda_col1_name", "เวลา")
                hdr_col2_name = data.get("agenda_col2_name", "รายละเอียด")
                hdr_cells[0].text = hdr_col1_name
                hdr_cells[1].text = hdr_col2_name
                set_cell_margins(hdr_cells[0], top=120, bottom=120, left=120, right=120)
                set_cell_margins(hdr_cells[1], top=120, bottom=120, left=180, right=180)

                for c_idx, cell in enumerate(hdr_cells):
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.35
                        if c_idx == 0:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.font.name = STRICT_FONT_NAME
                            r.font.size = Pt(10)
                            r.bold = True

                for item in data["agenda"]:
                    row_cells = table.add_row().cells
                    p_time = row_cells[0].paragraphs[0]
                    p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_time.paragraph_format.space_before = Pt(0)
                    p_time.paragraph_format.space_after = Pt(0)
                    p_time.paragraph_format.line_spacing = 1.35
                    r_t = p_time.add_run(item.get("time", ""))
                    r_t.font.name = STRICT_FONT_NAME
                    r_t.font.size = Pt(10)
                    set_cell_margins(row_cells[0], top=120, bottom=120, left=120, right=120)

                    p = row_cells[1].paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.35

                    module_title = item.get("module_title", "")
                    if module_title:
                        run_m = p.add_run(module_title)
                        run_m.font.name = STRICT_FONT_NAME
                        run_m.font.size = Pt(10)
                        run_m.bold = True

                    first_topic = True
                    for topic in item.get("topics", []):
                        if isinstance(topic, dict):
                            t_title = topic.get("title") or topic.get("text") or ""
                            if t_title:
                                if first_topic and not module_title:
                                    p_t = p
                                    first_topic = False
                                else:
                                    p_t = row_cells[1].add_paragraph()
                                p_t.paragraph_format.space_before = Pt(0)
                                p_t.paragraph_format.space_after = Pt(2)
                                p_t.paragraph_format.line_spacing = 1.35
                                p_t.paragraph_format.left_indent = Inches(0.12)
                                p_t.paragraph_format.first_line_indent = Inches(-0.12)

                                pPr_t = p_t._p.get_or_add_pPr()
                                numPr_t = OxmlElement('w:numPr')
                                ilvl_t = OxmlElement('w:ilvl')
                                ilvl_t.set(qn('w:val'), '0')
                                numId_t = OxmlElement('w:numId')
                                numId_t.set(qn('w:val'), '1')
                                numPr_t.append(ilvl_t)
                                numPr_t.append(numId_t)
                                pPr_t.append(numPr_t)

                                clean_t = clean_bullet_text(t_title)
                                run_t = p_t.add_run(clean_t)
                                run_t.font.name = STRICT_FONT_NAME
                                run_t.font.size = Pt(10)

                            sub_list = topic.get("sub_topics") or topic.get("sub_bullets") or []
                            for sub_item in sub_list:
                                p_sub = row_cells[1].add_paragraph()
                                p_sub.paragraph_format.space_before = Pt(0)
                                p_sub.paragraph_format.space_after = Pt(2)
                                p_sub.paragraph_format.line_spacing = 1.35
                                p_sub.paragraph_format.left_indent = Inches(0.28)
                                p_sub.paragraph_format.first_line_indent = Inches(-0.12)

                                pPr_sub = p_sub._p.get_or_add_pPr()
                                numPr_sub = OxmlElement('w:numPr')
                                ilvl_sub = OxmlElement('w:ilvl')
                                ilvl_sub.set(qn('w:val'), '0')
                                numId_sub = OxmlElement('w:numId')
                                numId_sub.set(qn('w:val'), '1')
                                numPr_sub.append(ilvl_sub)
                                numPr_sub.append(numId_sub)
                                pPr_sub.append(numPr_sub)

                                clean_sub = clean_bullet_text(sub_item)
                                run_sub = p_sub.add_run(clean_sub)
                                run_sub.font.name = STRICT_FONT_NAME
                                run_sub.font.size = Pt(10)

                        elif isinstance(topic, str):
                            if first_topic and not module_title:
                                p_t = p
                                first_topic = False
                            else:
                                p_t = row_cells[1].add_paragraph()
                            p_t.paragraph_format.space_before = Pt(0)
                            p_t.paragraph_format.space_after = Pt(2)
                            p_t.paragraph_format.line_spacing = 1.35

                            topic_str = str(topic or "")
                            is_sub = topic_str.startswith("  ") or topic_str.startswith("\t") or topic_str.startswith(" -") or topic_str.startswith(" •") or topic_str.startswith("- ")
                            p_t.paragraph_format.left_indent = Inches(0.28 if is_sub else 0.12)
                            p_t.paragraph_format.first_line_indent = Inches(-0.12)

                            pPr_t = p_t._p.get_or_add_pPr()
                            numPr_t = OxmlElement('w:numPr')
                            ilvl_t = OxmlElement('w:ilvl')
                            ilvl_t.set(qn('w:val'), '0')
                            numId_t = OxmlElement('w:numId')
                            numId_t.set(qn('w:val'), '1')
                            numPr_t.append(ilvl_t)
                            numPr_t.append(numId_t)
                            pPr_t.append(numPr_t)

                            clean_t = clean_bullet_text(topic)
                            run_t = p_t.add_run(clean_t)
                            run_t.font.name = STRICT_FONT_NAME
                            run_t.font.size = Pt(10)

                    if item.get("workshop"):
                        render_workshop_cell(row_cells[1], item.get("workshop"))

                    set_cell_margins(row_cells[1], top=120, bottom=120, left=180, right=180)

                apply_table_borders(table)
                set_table_col_widths(table, [1.55, 5.10])
            else:
                table = doc.add_table(rows=1, cols=1)
                try:
                    table.style = 'Table Grid'
                except Exception:
                    pass
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False

                hdr_cells = table.rows[0].cells
                hdr_col2_name = data.get("agenda_col2_name", "รายละเอียดเนื้อหาการอบรม")
                hdr_cells[0].text = hdr_col2_name
                set_cell_margins(hdr_cells[0], top=120, bottom=120, left=180, right=180)

                for p in hdr_cells[0].paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.35
                    for r in p.runs:
                        r.font.name = STRICT_FONT_NAME
                        r.font.size = Pt(10)
                        r.bold = True

                for item in data["agenda"]:
                    row_cells = table.add_row().cells
                    p = row_cells[0].paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.35

                    module_title = item.get("module_title", "")
                    if module_title:
                        run_m = p.add_run(module_title)
                        run_m.font.name = STRICT_FONT_NAME
                        run_m.font.size = Pt(10)
                        run_m.bold = True

                    first_topic = True
                    for topic in item.get("topics", []):
                        if isinstance(topic, dict):
                            t_title = topic.get("title") or topic.get("text") or ""
                            if t_title:
                                if first_topic and not module_title:
                                    p_t = p
                                    first_topic = False
                                else:
                                    p_t = row_cells[0].add_paragraph()
                                p_t.paragraph_format.space_before = Pt(0)
                                p_t.paragraph_format.space_after = Pt(2)
                                p_t.paragraph_format.line_spacing = 1.35
                                p_t.paragraph_format.left_indent = Inches(0.12)
                                p_t.paragraph_format.first_line_indent = Inches(-0.12)

                                run_t = add_smart_bullet(p_t, t_title, font_size=10, bold=True)

                            sub_list = topic.get("sub_topics") or topic.get("sub_bullets") or []
                            for sub_item in sub_list:
                                p_sub = row_cells[0].add_paragraph()
                                p_sub.paragraph_format.space_before = Pt(0)
                                p_sub.paragraph_format.space_after = Pt(2)
                                p_sub.paragraph_format.line_spacing = 1.35
                                p_sub.paragraph_format.left_indent = Inches(0.28)
                                p_sub.paragraph_format.first_line_indent = Inches(-0.12)

                                run_sub = add_smart_bullet(p_sub, sub_item, font_size=10)

                        elif isinstance(topic, str):
                            if first_topic and not module_title:
                                p_t = p
                                first_topic = False
                            else:
                                p_t = row_cells[0].add_paragraph()
                            p_t.paragraph_format.space_before = Pt(0)
                            p_t.paragraph_format.space_after = Pt(2)
                            p_t.paragraph_format.line_spacing = 1.35

                            topic_str = str(topic or "")
                            is_sub = topic_str.startswith("  ") or topic_str.startswith("\t") or topic_str.startswith(" -") or topic_str.startswith(" •") or topic_str.startswith("- ")
                            p_t.paragraph_format.left_indent = Inches(0.28 if is_sub else 0.12)
                            p_t.paragraph_format.first_line_indent = Inches(-0.12)

                            run_t = add_smart_bullet(p_t, topic, font_size=10)

                    if item.get("workshop"):
                        render_workshop_cell(row_cells[0], item.get("workshop"))

                    set_cell_margins(row_cells[0], top=120, bottom=120, left=180, right=180)

                apply_table_borders(table)
                set_table_col_widths(table, [6.65])

            add_blank_line(doc)

        elif sec_key == "learning_methods" and data.get("learning_methods"):
            add_heading(doc, data.get("learning_methods_title", "รูปแบบการอบรม"), font_size=12)
            if data.get("learning_methods_subhead"):
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(0)
                p_sub.paragraph_format.space_after = Pt(3)
                p_sub.paragraph_format.line_spacing = 1.35
                r_sub = p_sub.add_run(data["learning_methods_subhead"])
                r_sub.font.name = STRICT_FONT_NAME
                r_sub.font.size = Pt(10)
                r_sub.bold = False
            for item in data["learning_methods"]:
                add_bullet_p(doc, item)
            add_blank_line(doc)

        elif sec_key == "target_audience" and data.get("target_audience"):
            add_heading(doc, data.get("target_audience_title", "ผู้ที่เหมาะสมกับหลักสูตรนี้"), font_size=12)
            if data.get("target_audience_subhead"):
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(0)
                p_sub.paragraph_format.space_after = Pt(3)
                p_sub.paragraph_format.line_spacing = 1.35
                r_sub = p_sub.add_run(data["target_audience_subhead"])
                r_sub.font.name = STRICT_FONT_NAME
                r_sub.font.size = Pt(10)
                r_sub.bold = False
            for item in data["target_audience"]:
                add_bullet_p(doc, item)
            add_blank_line(doc)

        elif sec_key == "duration" and data.get("duration"):
            add_heading(doc, data.get("duration_title", "ระยะเวลาหลักสูตร"), font_size=12)
            for item in data["duration"]:
                add_bullet_p(doc, item)
            add_blank_line(doc)

        elif sec_key == "frameworks" and data.get("frameworks"):
            add_heading(doc, data.get("frameworks_title", "Framework ที่ใช้"), font_size=12)
            for item in data["frameworks"]:
                add_bullet_p(doc, item)
            add_blank_line(doc)

        elif sec_key == "workshop_activities" and data.get("workshop_activities"):
            add_heading(doc, "Workshop Activities", font_size=12)
            for ws in data["workshop_activities"]:
                p_head = doc.add_paragraph()
                p_head.paragraph_format.space_before = Pt(4)
                p_head.paragraph_format.space_after = Pt(1)
                p_head.paragraph_format.line_spacing = 1.35
                r = p_head.add_run(f"{ws.get('title', '')}")
                r.font.name = STRICT_FONT_NAME
                r.font.size = Pt(10)
                r.bold = True
                if ws.get("description"):
                    p_desc = doc.add_paragraph()
                    p_desc.paragraph_format.space_before = Pt(0)
                    p_desc.paragraph_format.space_after = Pt(4)
                    p_desc.paragraph_format.line_spacing = 1.35
                    r_d = p_desc.add_run(ws.get("description", ""))
                    r_d.font.name = STRICT_FONT_NAME
                    r_d.font.size = Pt(10)
                if ws.get("bullets"):
                    for b in ws.get("bullets"):
                        add_bullet_p(doc, b)
            add_blank_line(doc)

        elif sec_key == "teaching_style" and data.get("teaching_style"):
            add_heading(doc, data.get("teaching_style_title", "สไตล์การสอนของวิทยากร"), font_size=12)
            if isinstance(data["teaching_style"], dict):
                if data["teaching_style"].get("text"):
                    add_rationale_p(doc, data["teaching_style"]["text"])
                    add_blank_line(doc)
                if data["teaching_style"].get("bullets"):
                    for item in data["teaching_style"]["bullets"]:
                        add_bullet_p(doc, item)
                    add_blank_line(doc)
            elif isinstance(data["teaching_style"], list):
                for item in data["teaching_style"]:
                    add_bullet_p(doc, item)
                add_blank_line(doc)

        elif sec_key == "expected_outcomes" and data.get("expected_outcomes"):
            add_heading(doc, data.get("expected_outcomes_title", "สิ่งที่ผู้อบรมจะได้รับ"), font_size=12)
            if data.get("expected_outcomes_subhead"):
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(0)
                p_sub.paragraph_format.space_after = Pt(3)
                p_sub.paragraph_format.line_spacing = 1.35
                r_sub = p_sub.add_run(data["expected_outcomes_subhead"])
                r_sub.font.name = STRICT_FONT_NAME
                r_sub.font.size = Pt(10)
                r_sub.bold = False

            if isinstance(data["expected_outcomes"], str):
                add_rationale_p(doc, data["expected_outcomes"])
                add_blank_line(doc)
            elif isinstance(data["expected_outcomes"], list):
                for item in data["expected_outcomes"]:
                    if isinstance(item, dict):
                        t_title = item.get("title") or item.get("text") or ""
                        if t_title:
                            add_subhead_p(doc, t_title)
                        sub_list = item.get("sub_bullets") or item.get("bullets") or []
                        for sub in sub_list:
                            add_bullet_p(doc, sub, left_indent_in=0.5, hanging_in=0.25)
                    elif isinstance(item, str):
                        if is_intro_sentence(item):
                            add_subhead_p(doc, item)
                        else:
                            add_bullet_p(doc, item, left_indent_in=0.5, hanging_in=0.25)
                add_blank_line(doc)

        elif sec_key == "remarks" and data.get("remarks"):
            add_heading(doc, data.get("remarks_title", "หมายเหตุ"), font_size=12)
            for item in data["remarks"]:
                p_rmk = doc.add_paragraph()
                p_rmk.paragraph_format.space_before = Pt(0)
                p_rmk.paragraph_format.space_after = Pt(3)
                p_rmk.paragraph_format.line_spacing = 1.35
                r_rmk = p_rmk.add_run(item)
                r_rmk.font.name = STRICT_FONT_NAME
                r_rmk.font.size = Pt(10)

        elif sec_key == "cti_core_values" and data.get("cti_core_values"):
            add_heading(doc, data.get("cti_core_values_title", "การเชื่อมโยงหลักสูตรกับ CTI Core Values"), font_size=12)
            table = doc.add_table(rows=1, cols=3)
            try:
                table.style = 'Table Grid'
            except Exception:
                pass
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Core Value"
            hdr_cells[1].text = "สิ่งที่ต้องการสร้าง"
            hdr_cells[2].text = "สิ่งที่หลักสูตรพัฒนา"
            for c_idx, cell in enumerate(hdr_cells):
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.35
                    for r in p.runs:
                        r.font.name = STRICT_FONT_NAME
                        r.font.size = Pt(10)
                        r.bold = True

            for item in data["cti_core_values"]:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get("core_value", "")
                row_cells[1].text = item.get("target_behavior", "")
                row_cells[2].text = item.get("course_development", "")
                for c_idx, cell in enumerate(row_cells):
                    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.35
                        for r in p.runs:
                            r.font.name = STRICT_FONT_NAME
                            r.font.size = Pt(10)
                            if c_idx == 0:
                                r.bold = True

            apply_table_borders(table)
            set_table_col_widths(table, [1.50, 2.50, 2.65])
            add_blank_line(doc)

        elif sec_key == "learning_journey" and data.get("learning_journey"):
            add_heading(doc, data.get("learning_journey_title", "Learning Journey"), font_size=12)
            for item in data["learning_journey"]:
                p_lj = doc.add_paragraph()
                p_lj.paragraph_format.space_before = Pt(0)
                p_lj.paragraph_format.space_after = Pt(3)
                p_lj.paragraph_format.line_spacing = 1.35
                r_lj = p_lj.add_run(item)
                r_lj.font.name = STRICT_FONT_NAME
                r_lj.font.size = Pt(10)
            add_blank_line(doc)

        elif sec_key == "additional_sections" and data.get("additional_sections"):
            for sec in data["additional_sections"]:
                add_heading(doc, sec.get("title", "หัวข้ออื่นๆ"), font_size=12)
                if sec.get("description"):
                    add_rationale_p(doc, sec.get("description"))
                if sec.get("bullets"):
                    for b in sec.get("bullets"):
                        add_bullet_p(doc, b)
                add_blank_line(doc)

        elif sec_key == "followup_program" and data.get("followup_program"):
            add_heading(doc, data.get("followup_program_title", "Follow-up Program"), font_size=12)
            if data.get("followup_program_subhead"):
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(0)
                p_sub.paragraph_format.space_after = Pt(3)
                p_sub.paragraph_format.line_spacing = 1.35
                r_sub = p_sub.add_run(data["followup_program_subhead"])
                r_sub.font.name = STRICT_FONT_NAME
                r_sub.font.size = Pt(10)

            fp = data["followup_program"]
            if fp.get("participant"):
                p_part = doc.add_paragraph()
                p_part.paragraph_format.space_before = Pt(3)
                p_part.paragraph_format.space_after = Pt(2)
                p_part.paragraph_format.line_spacing = 1.35
                r_pt = p_part.add_run(fp["participant"].get("title", "ผู้เข้าอบรม"))
                r_pt.font.name = STRICT_FONT_NAME
                r_pt.font.size = Pt(10)
                r_pt.bold = True
                for b in fp["participant"].get("bullets", []):
                    add_bullet_p(doc, b)

            if fp.get("coaching_guide"):
                p_cg = doc.add_paragraph()
                p_cg.paragraph_format.space_before = Pt(4)
                p_cg.paragraph_format.space_after = Pt(2)
                p_cg.paragraph_format.line_spacing = 1.35
                r_cg = p_cg.add_run(fp["coaching_guide"].get("title", "HR / หัวหน้างาน"))
                r_cg.font.name = STRICT_FONT_NAME
                r_cg.font.size = Pt(10)
                r_cg.bold = True
                for q_item in fp["coaching_guide"].get("questions", []):
                    p_q = doc.add_paragraph()
                    p_q.paragraph_format.space_before = Pt(2)
                    p_q.paragraph_format.space_after = Pt(1)
                    p_q.paragraph_format.line_spacing = 1.35
                    r_q = p_q.add_run(q_item.get("q", ""))
                    r_q.font.name = STRICT_FONT_NAME
                    r_q.font.size = Pt(10)
                    r_q.bold = True
                    if q_item.get("desc"):
                        p_d = doc.add_paragraph()
                        p_d.paragraph_format.space_before = Pt(0)
                        p_d.paragraph_format.space_after = Pt(3)
                        p_d.paragraph_format.line_spacing = 1.35
                        p_d.paragraph_format.left_indent = Inches(0.25)
                        r_d = p_d.add_run(q_item["desc"])
                        r_d.font.name = STRICT_FONT_NAME
                        r_d.font.size = Pt(10)
                        r_d.bold = False
            add_blank_line(doc)

    # Enforce Sarabun 10pt on Normal style, empty paragraphs, and runs
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = STRICT_FONT_NAME
        normal_style.font.size = Pt(10)
    except Exception:
        pass

    for p in doc.paragraphs:
        if not p.runs:
            r = p.add_run("")
            r.font.name = STRICT_FONT_NAME
            r.font.size = Pt(10)
        else:
            for r in p.runs:
                r.font.name = STRICT_FONT_NAME
                if r.font.size is None:
                    r.font.size = Pt(10)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.runs:
                        r = p.add_run("")
                        r.font.name = STRICT_FONT_NAME
                        r.font.size = Pt(10)
                    else:
                        for r in p.runs:
                            r.font.name = STRICT_FONT_NAME
                            if r.font.size is None:
                                r.font.size = Pt(10)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    # Add debug info for model usage and fallback reason
    if data.get("_ai_model_used"):
        p_debug = doc.add_paragraph()
        p_debug.paragraph_format.space_before = Pt(24)
        p_debug.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        debug_text = f"[Generated by: {data.get('_ai_model_used')} | Keys Loaded: {data.get('_keys_loaded', 'Unknown')}]"
        if data.get("_fallback_reason"):
            debug_text += f"\n[Fallback Reason: {data.get('_fallback_reason')}]"
        r_debug = p_debug.add_run(debug_text)
        r_debug.font.name = "Arial"
        r_debug.font.size = Pt(8)
        r_debug.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(output_path)
    print(f"Successfully generated course outline at: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) > 2:
        json_path = sys.argv[1]
        out_path = sys.argv[2]
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        generate_doc(data, out_path)
    else:
        print("Usage: python generate_course_outline.py <data.json> <output.docx>")
